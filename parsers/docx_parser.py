from pathlib import Path
from typing import Optional, Dict, Any
import logging

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class FormattedParagraph:
    """表示带格式信息的段落"""
    def __init__(self, text: str, is_bold: bool = False, is_strike: bool = False, is_underline: bool = False):
        self.text = text
        self.is_bold = is_bold  # 加粗 = 重点
        self.is_strike = is_strike  # 删除线 = 删除这题
        self.is_underline = is_underline  # 下划线 = 要点


def _get_paragraph_number(para) -> str:
    """获取段落的编号（处理Word自动编号）"""
    try:
        # 检查段落是否有编号格式
        pPr = para._element.pPr
        if pPr is None:
            return ""
        
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return ""
        
        # 获取编号级别
        ilvl = numPr.find(qn('w:ilvl'))
        numId = numPr.find(qn('w:numId'))
        
        if ilvl is not None and numId is not None:
            # 这是一个编号段落，但无法直接获取显示的数字
            # 尝试从前面的段落推断编号
            return "[NUM]"  # 标记为编号段落
        
        return ""
    except:
        return ""


def parse_docx_file(path: str) -> Optional[str]:
    """解析DOCX文件，返回纯文本（尝试提取自动编号）"""
    file_path = Path(path)
    if not file_path.exists():
        logger.error("DOCX file not found: %s", path)
        return None
    try:
        doc = Document(str(file_path))
        lines = []
        num_counter = 0  # 编号计数器
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            text = para.text
            num_marker = _get_paragraph_number(para)
            
            # 如果检测到编号段落但文本中没有数字开头，添加编号
            if num_marker == "[NUM]" and not text.strip()[0].isdigit():
                num_counter += 1
                text = f"{num_counter}. {text}"
            elif text.strip()[0].isdigit():
                # 如果文本本身有数字开头，更新计数器
                import re
                match = re.match(r'^(\d+)', text.strip())
                if match:
                    num_counter = int(match.group(1))
            
            lines.append(text)
        
        return "\n".join(lines)
    except Exception as exc:
        logger.error("Failed to read DOCX file %s: %s", path, exc)
        return None


def parse_docx_file_with_format(path: str) -> Optional[list[FormattedParagraph]]:
    """解析DOCX文件，保留格式信息（加粗/删除线/下划线）"""
    file_path = Path(path)
    if not file_path.exists():
        logger.error("DOCX file not found: %s", path)
        return None
    try:
        doc = Document(str(file_path))
        formatted_paras = []
        num_counter = 0
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # 检测段落级别的格式
            is_para_bold = False
            is_para_strike = False
            is_para_underline = False
            
            # 遍历所有 runs (文本块)
            for run in para.runs:
                if run.bold:
                    is_para_bold = True
                if run.font.strike:
                    is_para_strike = True
                if run.underline:
                    is_para_underline = True
            
            text = para.text
            num_marker = _get_paragraph_number(para)
            
            # 处理自动编号
            if num_marker == "[NUM]" and not text.strip()[0].isdigit():
                num_counter += 1
                text = f"{num_counter}. {text}"
            elif text.strip()[0].isdigit():
                import re
                match = re.match(r'^(\d+)', text.strip())
                if match:
                    num_counter = int(match.group(1))
            
            formatted_para = FormattedParagraph(
                text=text,
                is_bold=is_para_bold,
                is_strike=is_para_strike,
                is_underline=is_para_underline
            )
            formatted_paras.append(formatted_para)
        
        return formatted_paras
    except Exception as exc:
        logger.error("Failed to read DOCX file %s: %s", path, exc)
        return None

